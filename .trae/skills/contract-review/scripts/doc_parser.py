"""
文档解析器
支持 PDF、Word (docx)、纯文本、图片 (OCR)

版本: 1.1.0
更新: 2025-02-04
"""

import os
import logging
from pathlib import Path
from typing import Optional, Union, Callable
from dataclasses import dataclass

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 可选依赖检测
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False


# 文件大小限制 (50MB)
MAX_FILE_SIZE = 50 * 1024 * 1024


@dataclass
class ParsedDocument:
    """解析后的文档"""
    filename: str
    file_type: str
    content: str
    page_count: int
    metadata: dict

    @property
    def char_count(self) -> int:
        return len(self.content)

    def preview(self, max_chars: int = 500) -> str:
        """返回内容预览"""
        if len(self.content) <= max_chars:
            return self.content
        return self.content[:max_chars] + "..."


class DocumentParser:
    """文档解析器"""

    # 支持的文件类型
    SUPPORTED_TYPES = {
        # 文本类
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "txt",
        ".md": "markdown",
        # 图片类 (需要 OCR)
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".gif": "image",
        ".webp": "image",
    }

    # 不支持的类型 (明确告知用户)
    UNSUPPORTED_TYPES = {
        ".doc": "旧版 Word 97-2003 格式，请转换为 .docx",
        ".xls": "Excel 格式，暂不支持",
        ".xlsx": "Excel 格式，暂不支持",
        ".ppt": "PowerPoint 格式，暂不支持",
        ".pptx": "PowerPoint 格式，暂不支持",
    }

    def __init__(self, ocr_callback: Optional[Callable[[Path], str]] = None):
        """
        初始化文档解析器

        Args:
            ocr_callback: OCR 回调函数，用于处理图片。
                          如果不提供，图片将无法解析。
                          函数签名: (image_path: Path) -> str
        """
        self.ocr_callback = ocr_callback
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        missing = []
        if not HAS_PDFPLUMBER:
            missing.append("pdfplumber")
        if not HAS_DOCX:
            missing.append("python-docx")
        if not HAS_CHARDET:
            missing.append("chardet (可选，用于编码检测)")

        if missing:
            logger.warning(f"部分功能需要安装依赖: pip install {' '.join(missing)}")

    def _check_file_size(self, file_path: Path):
        """检查文件大小"""
        size = file_path.stat().st_size
        if size > MAX_FILE_SIZE:
            raise ValueError(
                f"文件过大: {size / 1024 / 1024:.1f}MB，"
                f"最大支持 {MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
            )

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        解析文档

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument 对象

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件类型或文件过大
            ImportError: 缺少必要的依赖
        """
        file_path = Path(file_path)

        # 检查文件是否存在
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查文件大小
        self._check_file_size(file_path)

        suffix = file_path.suffix.lower()

        # 检查是否为不支持的类型
        if suffix in self.UNSUPPORTED_TYPES:
            raise ValueError(
                f"不支持的文件格式 {suffix}: {self.UNSUPPORTED_TYPES[suffix]}"
            )

        # 检查是否为支持的类型
        if suffix not in self.SUPPORTED_TYPES:
            raise ValueError(
                f"不支持的文件类型: {suffix}。"
                f"支持的格式: {', '.join(self.SUPPORTED_TYPES.keys())}"
            )

        file_type = self.SUPPORTED_TYPES[suffix]

        # 根据文件类型调用对应的解析器
        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type == "docx":
            return self._parse_docx(file_path)
        elif file_type in ("txt", "markdown"):
            return self._parse_text(file_path)
        elif file_type == "image":
            return self._parse_image(file_path)
        else:
            # 理论上不会到达这里，但为了类型安全
            raise ValueError(f"未知的文件类型: {file_type}")

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """解析 PDF 文件"""
        if not HAS_PDFPLUMBER:
            raise ImportError("解析 PDF 需要安装 pdfplumber: pip install pdfplumber")

        content_parts = []
        page_count = 0
        metadata = {}

        logger.info(f"解析 PDF: {file_path.name}")

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            metadata = dict(pdf.metadata) if pdf.metadata else {}

            for i, page in enumerate(pdf.pages):
                page_content = []

                # 提取文本
                text = page.extract_text()
                if text:
                    page_content.append(text)

                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        table_text = self._format_table(table)
                        if table_text:
                            page_content.append(f"\n[表格]\n{table_text}")

                if page_content:
                    content_parts.append(f"--- 第 {i + 1} 页 ---\n" + "\n".join(page_content))

        return ParsedDocument(
            filename=file_path.name,
            file_type="pdf",
            content="\n\n".join(content_parts),
            page_count=page_count,
            metadata=metadata,
        )

    def _format_table(self, table: list) -> str:
        """格式化表格为文本"""
        if not table:
            return ""

        rows = []
        for row in table:
            if row:
                cells = [str(cell).strip() if cell else "" for cell in row]
                rows.append(" | ".join(cells))

        return "\n".join(rows)

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """解析 Word 文档 (.docx)"""
        if not HAS_DOCX:
            raise ImportError("解析 Word 需要安装 python-docx: pip install python-docx")

        logger.info(f"解析 Word: {file_path.name}")

        doc = Document(file_path)
        content_parts = []

        # 提取元数据
        metadata = {}
        try:
            props = doc.core_properties
            metadata = {
                "author": props.author,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "title": props.title,
                "subject": props.subject,
            }
        except Exception as e:
            logger.warning(f"提取文档元数据失败: {e}")

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)

        # 提取表格
        for table in doc.tables:
            table_content = self._extract_docx_table(table)
            if table_content:
                content_parts.append(f"\n[表格]\n{table_content}")

        # 估算页数 (按每页约 500 字计算)
        total_chars = sum(len(p) for p in content_parts)
        estimated_pages = max(1, total_chars // 500)

        return ParsedDocument(
            filename=file_path.name,
            file_type="docx",
            content="\n\n".join(content_parts),
            page_count=estimated_pages,
            metadata=metadata,
        )

    def _extract_docx_table(self, table) -> str:
        """提取 Word 表格内容"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        return "\n".join(rows) if rows else ""

    def _parse_text(self, file_path: Path) -> ParsedDocument:
        """解析纯文本文件"""
        logger.info(f"解析文本: {file_path.name}")

        # 检测编码
        encoding = self._detect_encoding(file_path)
        logger.debug(f"检测到编码: {encoding}")

        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
        except UnicodeDecodeError:
            # 回退到 latin-1 (可以读取任何字节)
            logger.warning(f"使用 {encoding} 解码失败，尝试 latin-1")
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()

        return ParsedDocument(
            filename=file_path.name,
            file_type="txt",
            content=content,
            page_count=1,
            metadata={"encoding": encoding},
        )

    def _detect_encoding(self, file_path: Path) -> str:
        """检测文件编码"""
        if HAS_CHARDET:
            with open(file_path, "rb") as f:
                raw = f.read(10000)  # 读取前 10KB 用于检测
                result = chardet.detect(raw)
                if result and result.get("encoding"):
                    return result["encoding"]

        # 默认尝试顺序
        return "utf-8"

    def _parse_image(self, file_path: Path) -> ParsedDocument:
        """解析图片 (需要 OCR)"""
        if self.ocr_callback is None:
            raise ValueError(
                f"解析图片需要 OCR 能力。请在初始化时提供 ocr_callback，"
                f"或使用 llm_client.py --ocr {file_path}"
            )

        logger.info(f"OCR 识别图片: {file_path.name}")

        content = self.ocr_callback(file_path)

        return ParsedDocument(
            filename=file_path.name,
            file_type="image",
            content=content,
            page_count=1,
            metadata={"ocr": True},
        )


def parse_document(file_path: str, ocr_callback: Optional[Callable] = None) -> str:
    """便捷函数：解析文档并返回文本内容"""
    parser = DocumentParser(ocr_callback=ocr_callback)
    doc = parser.parse(file_path)
    return doc.content


def parse_with_info(file_path: str, ocr_callback: Optional[Callable] = None) -> dict:
    """便捷函数：解析文档并返回完整信息"""
    parser = DocumentParser(ocr_callback=ocr_callback)
    doc = parser.parse(file_path)
    return {
        "filename": doc.filename,
        "file_type": doc.file_type,
        "content": doc.content,
        "page_count": doc.page_count,
        "metadata": doc.metadata,
        "char_count": doc.char_count,
    }


# 命令行使用
if __name__ == "__main__":
    import sys

    def print_usage():
        print("用法: python doc_parser.py <文件路径> [--ocr]")
        print("")
        print("支持的格式:")
        print("  文本类: PDF, DOCX, TXT, MD")
        print("  图片类: JPG, PNG, GIF, WEBP (需要 --ocr 参数)")
        print("")
        print("选项:")
        print("  --ocr    对图片进行 OCR 识别 (需要配置 QWEN_API_KEY)")
        print("")
        print("示例:")
        print("  python doc_parser.py contract.pdf")
        print("  python doc_parser.py contract.jpg --ocr")

    if len(sys.argv) < 2:
        print_usage()
        sys.exit(1)

    file_path = sys.argv[1]
    use_ocr = "--ocr" in sys.argv

    # 如果需要 OCR，导入 llm_client
    ocr_callback = None
    if use_ocr:
        try:
            from llm_client import QwenVLClient
            vl_client = QwenVLClient()
            ocr_callback = lambda p: vl_client.ocr(p)
            logger.info("OCR 功能已启用 (使用 Qwen-VL)")
        except Exception as e:
            logger.error(f"初始化 OCR 失败: {e}")
            print("OCR 需要配置 QWEN_API_KEY 环境变量")
            sys.exit(1)

    try:
        info = parse_with_info(file_path, ocr_callback=ocr_callback)
        print(f"文件名: {info['filename']}")
        print(f"类型: {info['file_type']}")
        print(f"页数: {info['page_count']}")
        print(f"字符数: {info['char_count']}")
        print("-" * 50)
        print("内容预览 (前 2000 字):")
        print(info["content"][:2000])
        if len(info["content"]) > 2000:
            print(f"\n... (共 {info['char_count']} 字符)")
    except Exception as e:
        logger.error(f"解析失败: {e}")
        sys.exit(1)
