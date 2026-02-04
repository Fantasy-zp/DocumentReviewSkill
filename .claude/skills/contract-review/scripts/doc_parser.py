"""
文档解析器
支持 PDF、Word（docx）、纯文本
"""

import os
from pathlib import Path
from typing import Optional, Union
from dataclasses import dataclass

# 可选依赖
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@dataclass
class ParsedDocument:
    """解析后的文档"""
    filename: str
    file_type: str
    content: str
    page_count: int
    metadata: dict


class DocumentParser:
    """文档解析器"""

    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".txt": "txt",
        ".md": "markdown",
    }

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        missing = []
        if not HAS_PDFPLUMBER:
            missing.append("pdfplumber")
        if not HAS_DOCX:
            missing.append("python-docx")

        if missing:
            print(f"提示: 部分功能需要安装依赖: pip install {' '.join(missing)}")

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """解析文档"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件类型: {suffix}")

        file_type = self.SUPPORTED_TYPES[suffix]

        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return self._parse_docx(file_path)
        elif file_type in ("txt", "markdown"):
            return self._parse_text(file_path)

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """解析 PDF 文件"""
        if not HAS_PDFPLUMBER:
            raise ImportError("解析 PDF 需要安装 pdfplumber: pip install pdfplumber")

        content_parts = []
        page_count = 0
        metadata = {}

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            metadata = pdf.metadata or {}

            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    content_parts.append(f"--- 第 {i + 1} 页 ---\n{text}")

        return ParsedDocument(
            filename=file_path.name,
            file_type="pdf",
            content="\n\n".join(content_parts),
            page_count=page_count,
            metadata=metadata,
        )

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """解析 Word 文档"""
        if not HAS_DOCX:
            raise ImportError("解析 Word 需要安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        content_parts = []
        metadata = {
            "author": doc.core_properties.author,
            "created": str(doc.core_properties.created) if doc.core_properties.created else None,
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
        }

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)

        # 提取表格
        for table in doc.tables:
            table_content = self._extract_table(table)
            if table_content:
                content_parts.append(table_content)

        return ParsedDocument(
            filename=file_path.name,
            file_type="docx",
            content="\n\n".join(content_parts),
            page_count=len(doc.paragraphs) // 30 + 1,  # 估算页数
            metadata=metadata,
        )

    def _extract_table(self, table) -> str:
        """提取表格内容"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        if rows:
            return "\n".join(rows)
        return ""

    def _parse_text(self, file_path: Path) -> ParsedDocument:
        """解析纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return ParsedDocument(
            filename=file_path.name,
            file_type="txt",
            content=content,
            page_count=1,
            metadata={},
        )


def parse_document(file_path: str) -> str:
    """便捷函数：解析文档并返回文本内容"""
    parser = DocumentParser()
    doc = parser.parse(file_path)
    return doc.content


def parse_with_info(file_path: str) -> dict:
    """便捷函数：解析文档并返回完整信息"""
    parser = DocumentParser()
    doc = parser.parse(file_path)
    return {
        "filename": doc.filename,
        "file_type": doc.file_type,
        "content": doc.content,
        "page_count": doc.page_count,
        "metadata": doc.metadata,
        "char_count": len(doc.content),
    }


# 命令行使用
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法: python doc_parser.py <文件路径>")
        print("支持格式: PDF, DOCX, TXT, MD")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        info = parse_with_info(file_path)
        print(f"文件名: {info['filename']}")
        print(f"类型: {info['file_type']}")
        print(f"页数: {info['page_count']}")
        print(f"字符数: {info['char_count']}")
        print("-" * 50)
        print("内容预览 (前 2000 字):")
        print(info["content"][:2000])
    except Exception as e:
        print(f"解析失败: {e}")
        sys.exit(1)
